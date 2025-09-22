import docx

def parse_word(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]

        # Extract table contents
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        table_texts.append(cell.text.strip())

        text = "\n".join(paragraphs)

        # Include table contents
        if table_texts:
            text += "\n\n[Table Contents]\n" + "\n".join(table_texts)

        return text
    except Exception as e:
        # Handle encrypted files, permission issues, corrupted files, etc.
        return ""